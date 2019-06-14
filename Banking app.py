import docx
#if docx doesnt work use (pip3 install docx in terminal)

def produce_report():
    doc = docx.Document()
    doc.add_paragraph(serial_number)
    doc.add_paragraph(name)
    doc.add_paragraph(address)
    doc.add_paragraph(account_type)
    doc.save('data_report.docx')


print("XYZ Banking")

bank_data = {}
user_data = {}

while True:
    print("""
    What would you like to do?
    
    1. Create account
    2. Manage an existing account
    3. Exit
    4. Produce report of accounts
    """)
    option = int(input("Please input the number you wish to choose: "))

    if option == 1:
        name = input("Please enter your name: ")
        age = int(input("Please enter your age: "))
        address = input("What is your address?: ")
        account_type = input("Account type you want to open (savings or current): ")
        account_number = int(input("Please enter a 4 digit account number: "))
        money = int(input("Enter how much you want to deposit: "))

        bank_data[account_number] = user_data

        print("""
Your account has now been created, please log in.""")

    elif option == 2:
        account_number = int(input("Enter account no:"))
        if account_number in bank_data:
            print("Welcome back", name)

            while True:
                print("Home page")
                option2 = int(input("""
    1. Check balance
    2. Make a deposit
    3. Make a withdrawal            
    4. Exit to main menu
    Please choose an option: """))

                if option2 == 1:
                    print("Your balance is:", money)

                elif option2 == 2:
                    deposit = int(input("How much would you like to deposit: "))
                    money = money + deposit
                    print("Your new balance is: ", money)

                elif option2 == 3:
                    withdrawal = int(input("How much would you like to withdraw?: "))
                    money = money - withdrawal
                    print("You have withdrawn", withdrawal, "and your new balance is", money)

                elif option2 == 4:
                    break

                else:
                    print("Invalid option")

    elif option == 3:
        break

    elif option == 4:
        serial_number = name[0], account_type[0].upper(), "2019"
        produce_report()
        print("Report Created in folder")


    else:
        print("Invalid option")

print("Thank you for banking with XYZ Banking")
